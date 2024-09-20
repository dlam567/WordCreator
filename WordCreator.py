# noinspection SpellCheckingInspection
'''
    名称：     WordCreator.py
    修改日期：   2024-09-21
    版本：        1.0.0
    作者：       dlam567
'''
from docx import Document
import tkinter as tk
from tkinter import ttk, messagebox
import json, os

from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog
from docx.shared import Inches, Pt

# === 基本配置 ===

version = "1.0.0"

jsonPath = './settings.json'
wdTitle = ''
wdText = ''
wd_grade_class = '(请选择)'
textNumber = 0
title_font_size = 16
title_font_bold = True
title_font = "宋体"
text_font = "宋体"
text_font_size = 12
row_spacing = 1.5
left_page_spacing = 1
right_page_spacing = 1
top_page_spacing = 1
bottom_page_spacing = 1
school = "福州十中"
grade = ["高一","高二","高三"]
class_ = [1,2,3,4,5,6,7,8]
class_txt = None; grade_txt = None


# === 数据处理 ===
try:
    # 读取JSON文件
    with open('settings.json', 'r', encoding='utf-8') as f:
        settings = json.load(f)

    # 将JSON文件转换为Python变量
    title_font_size = settings["titleFontSize"]
    title_font_bold = settings["titleFontBold"]
    title_font = settings["titleFont"]
    text_font = settings["textFont"]
    text_font_size = settings["textFontSize"]
    row_spacing = settings["rowSpacing"]
    left_page_spacing = settings["leftPageSpacing[cm]"]
    right_page_spacing = settings["rightPageSpacing[cm]"]
    top_page_spacing = settings["topPageSpacing[cm]"]
    bottom_page_spacing = settings["bottomPageSpacing[cm]"]
    school = settings["school"]
    grade = settings["grade"]
    class_ = settings["class"]
except:
    print("读取配置文件失败")
    messagebox.showerror("错误", "读取配置文件失败")
    

def win_center(win_, width, height):
    screenwidth = win.winfo_screenwidth()
    screenheight = win.winfo_screenheight()
    size = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2 - 40)
    win_.geometry(size)

# === 班级选择窗口 ===
class_selection_window = None

def on_focus(event):
    event.widget.select_range(0, tk.END)

def open_class_selection_window():
    global class_selection_window, grade_txt, class_txt
    if class_selection_window is None or not class_selection_window.winfo_exists():
        class_selection_window = tk.Toplevel(win)
        class_selection_window.title("选择班级")
        class_selection_window.geometry("280x140")
        class_selection_window.resizable(False, False)  # 设置窗口大小不可调整
        class_selection_window.attributes("-topmost", True)  # 将窗口置于顶层

        # 获取年级和班级列表
        grade_list = grade
        class_list = class_

        # 创建下拉列表
        grade_var = tk.StringVar()
        if grade_txt:
            grade_var.set(grade_txt)
        else:
            grade_var.set("(请选择年级)")
        grade_option_menu = ttk.Combobox(class_selection_window, textvariable=grade_var, values=grade_list)
        grade_option_menu.pack(pady=10)
        grade_option_menu.bind("<FocusIn>",on_focus)

        class_var = tk.StringVar()
        if class_txt:
            class_var.set(class_txt)
        else:    
            class_var.set("(请选择班级)")
        class_option_menu = ttk.Combobox(class_selection_window, textvariable=class_var, values=class_list)
        class_option_menu.pack(pady=10)
        class_option_menu.bind("<FocusIn>",on_focus)

        # 确认按钮
        def confirm_class():
            global wd_grade_class,grade_txt,class_txt
            if grade_var.get() == "(请选择年级)" or class_var.get() == "(请选择班级)":
                messagebox.showerror("错误", "请选择年级和班级")
                class_selection_window.focus_set()  # 激活窗口
            else:
                grade_txt = grade_var.get()
                class_txt = class_var.get()
                wd_grade_class = f"{grade_txt} ({class_txt})班"
                textLabel_wdGradeAndClass.config(text=f"班级：{wd_grade_class}")
                class_selection_window.destroy()

        confirm_button = tk.Button(class_selection_window, text="确认", command=confirm_class)
        confirm_button.pack(pady=10)
        # 将窗口置于屏幕中央
        win_center(class_selection_window,280,140)
    else:
        class_selection_window.deiconify()  # 显示窗口
        # class_selection_window.lift()  # 将窗口置于顶层
        class_selection_window.attributes("-topmost", True)  # 将窗口置于顶层
        class_selection_window.focus_set()  # 激活窗口

# === 文本框内容变动检测 ===
def on_entry_changed(event):
    global textNumber, wdTitle, wdText
    # 获取文本总长度
    textNumber = len(text_wdTitle.get())+len(text_wdText.get("1.0", "end-1c"))
    textLabel_wdTextNumberCounter.config(text=f"总字数：{textNumber}")  # 刷新数据
    wdTitle = text_wdTitle.get()
    wdText = text_wdText.get("1.0", "end-1c")


# === 窗口主体 ===
win = tk.Tk()
win.title("统一格式Word生成器")
win_center(win,395,640)
if len(school)>=4:
    win_center(win,395+(len(school)-4)*16,640)
else:
    win_center(win,395,640)
win.resizable(False, False)  # 设置窗口大小不可调整
win.attributes("-topmost", True)  # 将窗口置于顶层
win.focus_set()  # 激活窗口


# === docx生成 ===
def create_docx():
    global doc
    if wd_grade_class == "(请选择)":
        messagebox.showerror("错误", "请选择年级和班级")
        win.focus_set()  # 激活窗口
    elif name_var.get() == "":
        messagebox.showerror("错误", "请填写姓名")
        win.focus_set()  # 激活窗口
    elif  wdTitle== "":
        messagebox.showerror("错误", "请填写文章标题")
        win.focus_set()  # 激活窗口
    else:
        try:
            doc = Document()
            doc.styles['Normal'].font.name = text_font
            from docx.oxml.ns import qn
            doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), text_font)


            # 添加一个标题
            title_doc = doc.add_paragraph(wdTitle)
            title_doc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_doc.runs[0].font.bold = bool(title_font_bold)


            # 添加一些段落
            name_doc = doc.add_paragraph(school+" "*4+wd_grade_class+" "+name_var.get())
            name_doc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(wdText)
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(left_page_spacing)
                section.right_margin = Inches(right_page_spacing)
                section.top_margin = Inches(top_page_spacing)
                section.bottom_margin = Inches(bottom_page_spacing)
            # 设置间距
            for paragraph in doc.paragraphs:
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = row_spacing
                for run in paragraph.runs:
                    run.font.name = text_font
                    run.font.size = Pt(text_font_size)
            # 针对标题
            title_doc.runs[0].font.name = title_font
            title_doc.runs[0].font.size = Pt(title_font_size)
        except:
            messagebox.showerror("错误", "生成失败，请检查输入内容")
        try:
            save_file_path = filedialog.asksaveasfilename(initialfile=f"{wdTitle}.docx",defaultextension=".docx", filetypes=[("Word文档", "*.docx")])
            # 保存文档
            doc.save(f"{save_file_path}")
            # messagebox.showinfo("提示", "生成成功")
            # 打开文档
            os.startfile(f"{save_file_path}")
        except:
            messagebox.showerror("错误", "保存失败，请检查文件路径")

# === 窗口部件 ===
textLabel_title = tk.Label(text="欢迎使用\n统一格式Word生成器",font=("微软雅黑",18,"bold"))
textLabel_title.grid(row=0, column=0, columnspan=3, padx=3, pady=3)

textLabel_wdGradeAndClass = tk.Label(text=f"班级：{wd_grade_class}",font=("微软雅黑"))
textLabel_wdGradeAndClass.grid(row=1, column=0,columnspan=3, padx=3, pady=3)

button_wdGradeAndClass = tk.Button(text="选择班级", height=1,command=open_class_selection_window)  # 这里接入新窗口选择
button_wdGradeAndClass.grid(row=1, column=2,columnspan=4, padx=3, pady=3)
name_var = tk.StringVar()
school_label = tk.Label(win, text=f"学校：{school}",font=("微软雅黑"))
school_label.grid(row=2, column=0, padx=3, pady=3)

name_label = tk.Label(win, text=f"姓名：",font=("微软雅黑"))
name_label.grid(row=2, column=1, padx=3, pady=3)

name_entry = tk.Entry(win, font=("微软雅黑"), textvariable=name_var)
name_entry.grid(row=2, column=2, padx=3, pady=3)

textLabel_wdTitle = tk.Label(text="文档标题",font=("微软雅黑"))
textLabel_wdTitle.grid(row=3, column=0, columnspan=3, padx=3, pady=3)

# 这里加入文本框
text_wdTitle = tk.Entry(font=("微软雅黑"),width=40)
text_wdTitle.grid(row=4, column=0, columnspan=3, padx=3, pady=3)

textLabel_wdText = tk.Label(text="正文",font=("微软雅黑"))
textLabel_wdText.grid(row=5, column=0, columnspan=3, padx=3, pady=3)

# 这里加入富文本框
scrollbar = tk.Scrollbar(win)
scrollbar.grid(row=6, column=5, sticky=tk.N+tk.S, padx=0, pady=3)

text_wdText = tk.Text(font=("微软雅黑"),width=40,height=7, yscrollcommand=scrollbar.set)
text_wdText.grid(row=6, column=0, columnspan=3, sticky='nsew',padx=3, pady=3)

scrollbar.config(command=text_wdText.yview)


text_wdTitle.bind('<KeyRelease>', on_entry_changed)
text_wdText.bind('<KeyRelease>', on_entry_changed)

textLabel_wdTextNumberCounter = tk.Label(text=f"总字数：{textNumber}",font=("微软雅黑"))
textLabel_wdTextNumberCounter.grid(row=7, column=0, columnspan=3, padx=3, pady=3)

button_wdGenerate = tk.Button(text="生成Word",font=("微软雅黑",14,"bold"),command=create_docx)  # 这里接入新窗口选择
button_wdGenerate.grid(row=8, column=0, columnspan=3, padx=3, pady=3)

textLabel_useTips = tk.Label(text="""使用方法:
1.将文章的标题输入“文档标题”对应文本框中
2.选择班级, 输入姓名
3.将文章正文粘贴入“正文”对应文本框中
4.点击“生成Word”按钮并选择存放文件的路径""",font=("微软雅黑",12,"bold"))
textLabel_useTips.grid(row=9, column=0, columnspan=3, padx=3, pady=3)

# button_settings = tk.Button(text="软件设置")  # 这里接入新窗口选择
# button_settings.grid(row=10, column=0, columnspan=3, padx=3, pady=3)

label_info = tk.Label(win, text=f"作者：dlam567   Version：{version}",font=("宋体",9))
label_info.grid(row=11, column=0, columnspan=3, padx=3, pady=13)

def open_github(event):
    import webbrowser
    webbrowser.open("https://github.com/dlam567")

label_info.bind('<Button-1>', open_github)

def right_menu(t):
    def callback1(event=None):
        global root
        t.event_generate('<<Cut>>')
        on_entry_changed(event)

    def callback2(event=None):
        global root
        t.event_generate('<<Copy>>')
        on_entry_changed(event)

    def callback3(event=None):
        global root
        t.event_generate('<<Paste>>')
        on_entry_changed(event)
    def callback4(event=None):
        global root
        t.event_generate('<<SelectAll>>')
        on_entry_changed(event)

    '''创建一个弹出菜单'''
    menu = tk.Menu(win,
                tearoff=False,
                #bg="black",
                )
    menu.add_command(label="剪切", command=callback1)
    menu.add_command(label="复制", command=callback2)
    menu.add_command(label="粘贴", command=callback3)
    menu.add_command(label="全选", command=callback4)

    def popup(event):
        menu.post(event.x_root, event.y_root)   # post在指定的位置显示弹出菜单

    t.bind("<Button-3>", popup)                 # 绑定鼠标右键,执行popup函数


right_menu(text_wdTitle)
right_menu(text_wdText)

win.mainloop()